from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

ACCENT = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
MUTED = RGBColor(0x4A, 0x55, 0x68)


def set_run_font(run, name="Calibri", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_hr(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A56DB")
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run_font(run, size=11, bold=True, color=ACCENT)
    add_hr(p)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    set_run_font(r, size=10, color=DARK)
    return p


def job_header(title, meta):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(meta)
    set_run_font(r2, size=9.5, italic=True, color=MUTED)


# HEADER
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_after = Pt(2)
r = name.add_run("EREN KÜÇÜK")
set_run_font(r, size=22, bold=True, color=DARK)

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.paragraph_format.space_before = Pt(0)
tag.paragraph_format.space_after = Pt(2)
r = tag.add_run("YBS Öğrencisi  ·  Yazılım & Kalite Mühendisliği  ·  Yapay Zeka / Prompt Mühendisliği")
set_run_font(r, size=10.5, bold=True, color=ACCENT)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_before = Pt(0)
contact.paragraph_format.space_after = Pt(2)
r = contact.add_run(
    "Düzce / Ürgüp, Türkiye  ·  +90 546 296 87 35  ·  ernkck81@gmail.com"
)
set_run_font(r, size=9.5, color=MUTED)

links = doc.add_paragraph()
links.alignment = WD_ALIGN_PARAGRAPH.CENTER
links.paragraph_format.space_before = Pt(0)
links.paragraph_format.space_after = Pt(2)
r = links.add_run(
    "Canlı ürün: https://qualiguard.com.tr    ·    GitHub: https://github.com/erenkkc/QualiGuard"
)
set_run_font(r, size=9.5, bold=True, color=ACCENT)

# SUMMARY
section_heading("Profesyonel Özet")
summary = doc.add_paragraph()
summary.paragraph_format.space_after = Pt(2)
txt = (
    "Kapadokya Üniversitesi Yönetim Bilişim Sistemleri öğrencisi; full-stack geliştirme, yazılım kalite "
    "mühendisliği ve yapay zeka destekli iş akışlarını bir arada yürüten bir profil. SonarQube benzeri "
    "statik analiz platformu QualiGuard’ı sıfırdan geliştirdi ve https://qualiguard.com.tr adresinde "
    "HTTPS ile canlıya aldı (Go, kalite kapısı, CI/CD, Docker, VS Code eklentisi, yerel LLM / Ollama). "
    "Prompt mühendisliğiyle sistem prompt’ları tasarlayıp "
    "halüsinasyonu azaltmaya, geliştiriciler için güvenilir Türkçe YZ asistanı üretmeye odaklanıyor. "
    "PHP/MySQL web projeleri, Cisco ağ temelleri ve kalite kontrol disiplinini birleştirerek staj veya "
    "junior roller için yazılım geliştirme, DevOps araçları, QA ya da yapay zeka ürün ekiplerine hazır."
)
r = summary.add_run(txt)
set_run_font(r, size=10, color=DARK)

# EXPERIENCE
section_heading("İş Deneyimi")
job_header(
    "Kalite Kontrol Görevlisi & Teknik Çizim Uzmanı (AutoCAD)",
    "Üretim / İmalat Ortamı",
)
bullet(
    "Üretim süreçlerinde uçtan uca kalite kontrol noktalarını yönetti; hataları erken yakalayarak yeniden iş riskini azalttı."
)
bullet(
    "Tasarım, planlama ve saha uygulamasını destekleyen AutoCAD teknik çizimlerini üretti ve güncelledi."
)
bullet(
    "İzlenebilirlik ve sürekli iyileştirme için denetlenebilir dokümantasyon ve muayene kayıtlarını tuttu."
)
bullet(
    "Teknik bulguları net iletişime dönüştürerek saha ile süreç standartları arasında köprü kurdu."
)

# QUALIGUARD
section_heading("Öne Çıkan Proje — QualiGuard")
job_header(
    "QualiGuard — Statik Kod Analizi ve Kalite Kapısı Platformu",
    "Canlı: https://qualiguard.com.tr  ·  GitHub: https://github.com/erenkkc/QualiGuard  ·  Go, Python, JS, REST, SQLite, Docker, Ollama, VS Code, GitHub Actions",
)
bullet(
    "SonarQube benzeri kod kalitesi platformunu sıfırdan tasarlayıp qualiguard.com.tr üzerinde HTTPS ile canlıya aldı (VPS, Docker, Caddy, Cloudflare DNS)."
)
bullet(
    "Ürün yüzeyi: CLI tarayıcı, analiz sunucusu, landing page, şifreli panel, white-label, canlı analiz, zip yükleme ve Ollama ile Türkçe YZ sohbet."
)
bullet(
    "Çok dilli statik analiz (Python, JavaScript/TypeScript, Go, Java, C#) ile güvenlik kuralları (SQL injection, eval, secret) ve stil denetimi (Ruff, ESLint) uyguladı."
)
bullet(
    "Kalite kapıları (GEÇER / UYARI / KALIR), sorun parmak izi, yanlış alarm bastırma, SARIF/JSON/HTML raporlama ve CI için PR yorumları geliştirdi."
)
bullet(
    "Geliştirici araçları teslim etti: VS Code / Cursor eklentisi (diagnostics, workspace tarama, hover açıklama) ve otomatik PR kontrolü için GitHub Actions iş akışı."
)
bullet(
    "Canlı operasyonu üstlendi: Natro VPS, Cloudflare nameserver, Caddy ile Let’s Encrypt, panel şifresi ve Google Search Console index kurulumu."
)

# OTHER PROJECTS
section_heading("Akademik ve Kişisel Projeler")
job_header(
    "Kütüphane Ödünç Takip Sistemi — YAZ307 Web Programlama II",
    "PHP, MySQL, PDO, Bootstrap, HTML/CSS",
)
bullet(
    "Ödünç alma, iade ve gecikme takibi için oturum tabanlı yönetici girişi olan full-stack üniversite kütüphane sistemi geliştirdi."
)
bullet(
    "PDO prepared statement ve XSS korumasıyla güvenliği güçlendirdi; PHP DateTime ile otomatik gecikme uyarısı ekledi."
)

job_header("Kullanıcı Yönetim Sistemi", "PHP, MySQL, PDO, HTML/CSS")
bullet(
    "Kayıt / giriş / profil / çıkış akışlı güvenli üyelik platformu kurdu; parola hash’leme, SQL injection önleme ve e-posta tekilliği uyguladı."
)

job_header("Mobil Dokunmatik Oyun", "HTML5 Canvas, JavaScript, CSS")
bullet(
    "Dokunmatik kontroller, can/skor sistemi, parçacık efektleri ve dinamik zorluk içeren responsive tarayıcı oyunu geliştirdi — yalnızca vanilla JS."
)

# AI / PROMPT
section_heading("Yapay Zeka ve Prompt Mühendisliği")
bullet(
    "Prompt mühendisliği odağı: güvenilir Türkçe geliştirici asistanları için sistem prompt’ları, few-shot örnekler ve çıktı kısıtları tasarlama."
)
bullet(
    "QualiGuard’a yerel LLM (Ollama / llama3.2) entegre ederek sorun açıklama ve serbest kod kalitesi sohbeti sağladı — verinin yerel kalmasına öncelik verdi."
)
bullet(
    "Halüsinasyonu azaltmak, İngilizce sızıntıyı filtrelemek ve yanıtları kural motoru ile kalite kapısı mantığına dayandırmak için prompt’ları iteratif iyileştirdi."
)
bullet(
    "Prompt tasarımı + statik analizin geliştirici oryantasyonu, kod incelemesi ve güvenlik farkındalığını nasıl hızlandırabileceğini araştırıyor."
)

# EDUCATION
section_heading("Eğitim")
job_header(
    "Kapadokya Üniversitesi — Ürgüp, Türkiye",
    "Lisans · Yönetim Bilişim Sistemleri (YBS)  ·  2023 – Devam",
)
bullet(
    "İlgili dersler: Bilgi Sistemleri, Veritabanı Yönetimi, İş Analitiği, BT Altyapısı."
)
bullet(
    "İş süreçlerini modern yazılım araçları, yapay zeka asistanları ve kalite otomasyonuyla birleştirmeye odaklanıyor."
)

# CERTS
section_heading("Sertifikalar")
bullet(
    "Cisco — Introduction to Packet Tracer: ağ simülasyonu, topoloji tasarımı, temel sorun giderme."
)
bullet(
    "Cisco — IT Essentials: bilgisayar donanımı, işletim sistemleri, ağ ve güvenlik temelleri."
)

# SKILLS
section_heading("Teknik Beceriler")
pairs = [
    (
        "Diller ve Web: ",
        "Go, PHP, Python, JavaScript/TypeScript, HTML5, CSS3, MySQL, Bootstrap, HTML5 Canvas",
    ),
    (
        "Kalite ve DevTools: ",
        "Statik analiz, kalite kapısı, SARIF, Ruff, ESLint, GitHub Actions, Docker, Caddy",
    ),
    (
        "YZ Yığını: ",
        "Prompt mühendisliği, Ollama, yerel LLM, yapay zeka destekli kod açıklama ve geliştirici sohbeti",
    ),
    (
        "Platform ve Araçlar: ",
        "VS Code / Cursor eklentileri, REST API, SQLite, AutoCAD, Packet Tracer, phpMyAdmin, Microsoft Office",
    ),
    (
        "Güçlü Yanlar: ",
        "Ürün düşüncesi, teknik dokümantasyon, kalite bakış açısı, sistem tasarımı, hızlı prototipleme, net iletişim",
    ),
]
for label, rest in pairs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(label)
    set_run_font(r1, size=10, bold=True, color=DARK)
    r2 = p.add_run(rest)
    set_run_font(r2, size=10, color=DARK)

# LANGUAGES
section_heading("Diller")
p = doc.add_paragraph()
r = p.add_run(
    "Türkçe: Ana dil    ·    İngilizce: Orta (B1+) — teknik okuma, profesyonel yazışma, İngilizce ve Türkçe prompt iterasyonu"
)
set_run_font(r, size=10, color=DARK)

# INTERESTS
section_heading("İlgi Alanları")
p = doc.add_paragraph()
r = p.add_run(
    "Yazılım kalite platformları · Prompt mühendisliği ve uygulamalı yapay zeka · Full-stack / backend geliştirme · "
    "Geliştirici deneyimi (DevTools) · CI/CD otomasyonu · Güvenli kodlama · BT altyapısı ve iş–teknoloji uyumu"
)
set_run_font(r, size=10, color=DARK)

paths = [
    r"c:\Users\Eren\OneDrive - kapadokya.edu.tr\Masaüstü\Ozgecmis_Eren_Kucuk_TR.docx",
    r"C:\Users\Eren\Desktop\QualiGuard\Ozgecmis_Eren_Kucuk_TR.docx",
    r"C:\Users\Eren\Desktop\Ozgecmis_Eren_Kucuk_TR.docx",
]
for path in paths:
    try:
        doc.save(path)
        print("SAVED", path)
    except Exception as e:
        print("FAIL", path, "->", e)
