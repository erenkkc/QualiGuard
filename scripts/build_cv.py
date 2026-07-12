from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

ACCENT = RGBColor(0x1A, 0x56, 0xDB)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
MUTED = RGBColor(0x4A, 0x55, 0x68)


def set_run_font(run, name="Calibri", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_hr(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A56DB")
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run_font(run, size=11, bold=True, color=ACCENT)
    add_hr(p)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    set_run_font(r, size=10, color=DARK)
    return p


def job_header(title, meta):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(meta)
    set_run_font(r2, size=9.5, italic=True, color=MUTED)


# HEADER
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_after = Pt(2)
r = name.add_run("EREN KÜÇÜK")
set_run_font(r, size=22, bold=True, color=DARK)

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.paragraph_format.space_before = Pt(0)
tag.paragraph_format.space_after = Pt(2)
r = tag.add_run("MIS Student  ·  Business & Technology  ·  Software Quality  ·  AI / Prompt Engineering")
set_run_font(r, size=10.5, bold=True, color=ACCENT)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_before = Pt(0)
contact.paragraph_format.space_after = Pt(2)
r = contact.add_run("Düzce / Ürgüp, Turkey  ·  +90 546 296 87 35  ·  ernkck81@gmail.com")
set_run_font(r, size=9.5, color=MUTED)

# SUMMARY
section_heading("Professional Summary")
summary = doc.add_paragraph()
summary.paragraph_format.space_after = Pt(2)
txt = (
    "Driven Management Information Systems student at Kapadokya University with advanced business education "
    "and a rare dual fluency in strategy and technology. Trained at a high academic level in management, "
    "organizational processes, business analytics, and decision-making — then applies that business mindset "
    "to build real software products. Creator of QualiGuard, a SonarQube-inspired static analysis platform "
    "(Go, quality gates, CI/CD, Docker, VS Code extension, and local LLM / Ollama). Practices prompt engineering "
    "to design reliable system prompts, reduce hallucination, and deliver Turkish-language AI assistance for "
    "developers. Combines rigorous business foundations, PHP/MySQL web projects, Cisco networking, and "
    "quality-control discipline into a profile ready for internship or junior roles in software, product, "
    "QA, DevOps-adjacent tooling, or AI-enabled business teams."
)
r = summary.add_run(txt)
set_run_font(r, size=10, color=DARK)

# EXPERIENCE
section_heading("Work Experience")
job_header(
    "Quality Control Officer & Technical Drafting Specialist (AutoCAD)",
    "Production / Manufacturing Environment",
)
bullet(
    "Owned end-to-end quality checkpoints across production workflows — catching defects early and cutting rework risk through rigorous inspection discipline."
)
bullet(
    "Produced and iterated AutoCAD technical drawings that directly supported design, planning, and on-floor implementation."
)
bullet(
    "Maintained audit-ready documentation and inspection records, strengthening traceability and continuous improvement culture."
)
bullet(
    "Translated technical findings into clear communication for cross-functional teams — bridging shop-floor reality with process standards."
)

# QUALIGUARD
section_heading("Flagship Project — QualiGuard")
job_header(
    "QualiGuard — Static Code Analysis & Quality Gate Platform",
    "Independent Product Build  ·  Go, Python, JavaScript, REST API, SQLite, Docker, Ollama, VS Code Extension, GitHub Actions",
)
bullet(
    "Architected and shipped a SonarQube-like code quality platform from scratch: CLI scanner, analysis server, web dashboard, and quality-gate engine."
)
bullet(
    "Implemented multi-language static analysis (Python, JavaScript/TypeScript, Go, Java, C#) with security rules (SQL injection, eval, secrets) and style linting (Ruff, ESLint)."
)
bullet(
    "Designed quality gates (PASS / WARN / FAIL) with issue fingerprinting, false-positive suppression, SARIF/JSON/HTML export, and PR comment decoration for CI."
)
bullet(
    "Built a production-ready web product surface: landing page, authenticated panel, white-label branding, live playground, zip upload, and Turkish AI chat via Ollama."
)
bullet(
    "Delivered developer tooling: VS Code / Cursor extension (diagnostics, workspace scan, hover explanations) and GitHub Actions workflow for automated PR quality checks."
)
bullet(
    "Engineered Docker + Caddy HTTPS deploy packaging and environment-based panel security — positioning QualiGuard as a shippable, domain-ready product."
)

# OTHER PROJECTS
section_heading("Academic & Personal Projects")
job_header(
    "Library Loan Tracking System — YAZ307 Web Programming II",
    "PHP, MySQL, PDO, Bootstrap, HTML/CSS",
)
bullet(
    "Delivered a full-stack university library system for loans, returns, and overdue tracking with session-based admin auth."
)
bullet(
    "Hardened security with PDO prepared statements and XSS protections; automated overdue detection with PHP DateTime alerts."
)

job_header("User Management System", "PHP, MySQL, PDO, HTML/CSS")
bullet(
    "Built a secure membership platform (register / login / profile / logout) with password hashing, SQL injection defense, and email uniqueness enforcement."
)

job_header("Mobile Tap Game", "HTML5 Canvas, JavaScript, CSS")
bullet(
    "Shipped a responsive browser game with touch controls, lives/score systems, particle effects, and dynamic difficulty — pure vanilla JS."
)

# AI / PROMPT
section_heading("AI & Prompt Engineering")
bullet(
    "Prompt engineering focus: crafting system prompts, few-shot patterns, and output constraints for reliable Turkish developer assistants."
)
bullet(
    "Integrated local LLMs (Ollama / llama3.2) into QualiGuard for issue explanation and free-form code-quality chat — prioritizing privacy (data stays local)."
)
bullet(
    "Iteratively refined prompts to reduce hallucination, filter English leakage, and keep answers grounded in QualiGuard’s rule engine and quality-gate semantics."
)
bullet(
    "Exploring how prompt design + static analysis can accelerate developer onboarding, code review, and security awareness."
)

# EDUCATION
section_heading("Education")
job_header(
    "Kapadokya University — Ürgüp, Turkey",
    "Bachelor of Science · Management Information Systems (MIS)  ·  2023 – Present",
)
bullet(
    "Completed advanced business education at a high academic level — covering management principles, organizational behavior, strategic decision-making, and enterprise process design."
)
bullet(
    "Core business & analytics foundation: Business Administration concepts, Business Analytics, information-driven decision support, operations thinking, and IT–business alignment."
)
bullet(
    "Relevant MIS coursework: Information Systems, Database Management, Business Analytics, IT Infrastructure — translating business requirements into technology solutions."
)
bullet(
    "Combines elite business literacy with hands-on engineering: builds products that solve real organizational problems, not just code for its own sake."
)

# CERTS
section_heading("Certifications")
bullet(
    "Cisco — Introduction to Packet Tracer: network simulation, topology design, basic troubleshooting."
)
bullet(
    "Cisco — IT Essentials: PC hardware, OS fundamentals, networking, and security basics."
)

# SKILLS
section_heading("Technical Skills")
pairs = [
    (
        "Languages & Web: ",
        "Go, PHP, Python, JavaScript/TypeScript, HTML5, CSS3, MySQL, Bootstrap, HTML5 Canvas",
    ),
    (
        "Quality & DevTools: ",
        "Static analysis, quality gates, SARIF, Ruff, ESLint, GitHub Actions, Docker, Caddy",
    ),
    (
        "AI Stack: ",
        "Prompt engineering, Ollama, local LLMs, AI-assisted code explanation & developer chat",
    ),
    (
        "Platforms & Tools: ",
        "VS Code / Cursor extensions, REST APIs, SQLite, AutoCAD, Packet Tracer, phpMyAdmin, Microsoft Office",
    ),
    (
        "Strengths: ",
        "Advanced business education, product thinking, stakeholder communication, QA mindset, systems design, rapid prototyping",
    ),
]
for label, rest in pairs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(label)
    set_run_font(r1, size=10, bold=True, color=DARK)
    r2 = p.add_run(rest)
    set_run_font(r2, size=10, color=DARK)

# LANGUAGES
section_heading("Languages")
p = doc.add_paragraph()
r = p.add_run(
    "Turkish: Native    ·    English: Intermediate (B1+) — technical reading, professional correspondence, AI prompt iteration in English & Turkish"
)
set_run_font(r, size=10, color=DARK)

# INTERESTS
section_heading("Areas of Interest")
p = doc.add_paragraph()
r = p.add_run(
    "Business strategy & process excellence · Software quality platforms · Prompt engineering & applied AI · "
    "Full-stack / backend development · DevTools & developer experience · CI/CD automation · "
    "Secure coding · IT–business transformation"
)
set_run_font(r, size=10, color=DARK)

paths = [
    r"c:\Users\Eren\OneDrive - kapadokya.edu.tr\Masaüstü\Resume_Eren_Kucuk_EN.docx",
    r"C:\Users\Eren\Desktop\QualiGuard\Resume_Eren_Kucuk_EN.docx",
    r"C:\Users\Eren\Desktop\Resume_Eren_Kucuk_EN.docx",
]
for path in paths:
    try:
        doc.save(path)
        print("SAVED", path)
    except Exception as e:
        print("FAIL", path, "->", e)
